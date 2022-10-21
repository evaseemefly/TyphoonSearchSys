# 作为 word中的段落
import abc
import os
from docx import Document

class IDocument(abc.ABC):
    # _document=None

    def __init__(self,dir,filename):
        self.dir=dir
        self.filename=filename
        self._document=None
        self._wordText=None
        fullpath=os.path.join(dir,filename)
        self.read(fullpath)

    @abc.abstractmethod
    def load(self,iterable):
        '''

        :param iterable:
        :return:
        '''

    @abc.abstractmethod
    def read(self,fullpath:str):
        '''
            读取文档
        :param fullpath:
        :return:
        '''


class Paragraph(IDocument):
    '''
        段落
    '''

    def __init__(self,dir:str,filename:str):
        # super(Paragraph, self).__init__(dir,filename)
        super().__init__(dir, filename)

    def load(self,iterable):
        pass

    def read(self,fullpath:str):
        '''
             根据全路径读取word文件并写入_wordText
        :return:
        '''
        fullWordText=[]
        if self._document is None:
            self._document=Document(fullpath)
            for para in self._document.paragraphs:
                if len(para.text) > 0:
                    fullWordText.append(para.text)
            self._wordText='\n'.join(fullWordText)

    @property
    def wordText(self):
        '''
            读取的word的text内容
        :return:
        '''
        return self._wordText

